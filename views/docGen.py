import os
import re
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from views.nvidia_llm import invoke, build_messages

_PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
load_dotenv(os.path.join(_PROJECT_ROOT, ".env"), override=True)

DOCUMENT_GENERATION_PROMPT = """\
You are an expert legal document drafter. Generate a complete, professional, ready-to-use legal document based on the user's request.

**Instructions:**
1. Identify the type of legal document requested (e.g., NDA, lease agreement, will, affidavit, complaint, contract, MOU, power of attorney, bail application, legal notice, partnership deed, employment agreement, divorce petition, sale deed, gift deed, promissory note, indemnity bond, cease and desist letter, or ANY other legal document).
2. Extract ALL relevant details from the user's prompt (names, dates, addresses, amounts, terms, conditions, etc.).
3. For any details NOT provided, use realistic placeholder values clearly marked with square brackets like [PLACEHOLDER].
4. Generate a COMPLETE, professionally formatted legal document with:
   - Proper title and header
   - Date and reference numbers
   - All necessary legal clauses and sections
   - Numbered articles/sections
   - Definitions section where appropriate
   - Governing law and jurisdiction clause
   - Dispute resolution clause
   - Signature blocks with witness sections
   - Any schedules or annexures if needed
5. Use proper legal language and terminology.
6. Follow Indian legal standards and conventions unless another jurisdiction is specified.
7. Make the document comprehensive - not a summary or outline, but a REAL usable document.

**OUTPUT FORMAT:**
Use markdown formatting with:
- # for document title
- ## for major sections
- ### for subsections
- **bold** for defined terms and emphasis
- Numbered lists for clauses
- --- for section separators
"""

def generate_legal_document(prompt, save_dir='static/generated_docs'):
    """
    Generates a complete legal document using NVIDIA AI and saves as Word .docx.
    
    Args:
        prompt (str): User's natural language description of the document they need.
        save_dir (str): Directory where the document will be saved.
    
    Returns:
        (str, str, str): file_path, file_name, preview_markdown
    """
    if not os.path.exists(save_dir):
        os.makedirs(save_dir)

    # Generate document content using NVIDIA API
    messages = build_messages(DOCUMENT_GENERATION_PROMPT, user_message=prompt)
    document_markdown = invoke(messages, temperature=0.7, max_tokens=4096)

    # Determine file name from document type
    file_name = _extract_filename(prompt, document_markdown)

    # Convert markdown to Word document
    doc = _create_docx_from_markdown(document_markdown)

    file_path = os.path.join(save_dir, file_name)
    doc.save(file_path)

    return file_path, file_name, document_markdown


def _extract_filename(prompt, content):
    """Extract a meaningful filename from the prompt or content."""
    prompt_lower = prompt.lower()
    
    doc_types = {
        "nda": "non_disclosure_agreement",
        "non-disclosure": "non_disclosure_agreement",
        "non disclosure": "non_disclosure_agreement",
        "lease": "lease_agreement",
        "rental": "rental_agreement",
        "employment": "employment_agreement",
        "contract": "contract",
        "will": "last_will_and_testament",
        "testament": "last_will_and_testament",
        "affidavit": "affidavit",
        "complaint": "legal_complaint",
        "mou": "memorandum_of_understanding",
        "memorandum": "memorandum_of_understanding",
        "power of attorney": "power_of_attorney",
        "bail": "bail_application",
        "notice": "legal_notice",
        "cease and desist": "cease_and_desist",
        "partnership": "partnership_deed",
        "divorce": "divorce_petition",
        "sale deed": "sale_deed",
        "gift deed": "gift_deed",
        "promissory": "promissory_note",
        "indemnity": "indemnity_bond",
        "service": "service_agreement",
        "freelance": "freelance_agreement",
        "consulting": "consulting_agreement",
        "loan": "loan_agreement",
        "settlement": "settlement_agreement",
        "petition": "legal_petition",
        "declaration": "declaration",
        "undertaking": "undertaking",
        "warranty": "warranty_agreement",
        "license": "license_agreement",
    }
    
    for keyword, name in doc_types.items():
        if keyword in prompt_lower:
            return f"{name}.docx"
    
    clean = re.sub(r'[^a-zA-Z\s]', '', prompt)
    words = clean.strip().split()[:4]
    name = "_".join(w.lower() for w in words) if words else "legal_document"
    return f"{name}.docx"


def _create_docx_from_markdown(markdown_text):
    """Convert markdown content to a professionally formatted Word document."""
    doc = DocxDocument()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(30, 30, 30)
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    lines = markdown_text.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        if line.startswith('# ') and not line.startswith('## '):
            title = _clean_markdown(line[2:])
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title.upper())
            run.bold = True
            run.font.size = Pt(16)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(20, 20, 80)
            p.space_after = Pt(12)
        
        elif line.startswith('## '):
            heading_text = _clean_markdown(line[3:])
            p = doc.add_paragraph()
            run = p.add_run(heading_text)
            run.bold = True
            run.font.size = Pt(13)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(30, 30, 80)
            p.space_before = Pt(12)
            p.space_after = Pt(6)
        
        elif line.startswith('### '):
            heading_text = _clean_markdown(line[4:])
            p = doc.add_paragraph()
            run = p.add_run(heading_text)
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            p.space_before = Pt(8)
            p.space_after = Pt(4)
        
        elif line.startswith('---'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('─' * 60)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(150, 150, 150)
        
        elif re.match(r'^\d+[\.)\]]\s', line) or line.startswith('- ') or line.startswith('* '):
            text = re.sub(r'^\d+[\.)\]]\s*', '', line)
            text = re.sub(r'^[-\*]\s*', '', text)
            text = _clean_markdown(text)
            p = doc.add_paragraph(style='List Number' if re.match(r'^\d+[\.)\]]\s', line) else 'List Bullet')
            _add_formatted_text(p, text)
        
        else:
            text = _clean_markdown(line)
            p = doc.add_paragraph()
            _add_formatted_text(p, text)
        
        i += 1
    
    return doc


def _clean_markdown(text):
    """Remove markdown artifacts but preserve structure for formatting."""
    text = re.sub(r'!\[.*?\]\(.*?\)', '', text)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    return text.strip()


def _add_formatted_text(paragraph, text):
    """Add text to a paragraph, handling bold and italic markdown."""
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        else:
            italic_parts = re.split(r'(\*[^*]+\*)', part)
            for ip in italic_parts:
                if ip.startswith('*') and ip.endswith('*') and not ip.startswith('**'):
                    run = paragraph.add_run(ip[1:-1])
                    run.italic = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                else:
                    run = paragraph.add_run(ip)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
